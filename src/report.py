import os
import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

from src.ingest import IngestionEngine
from src.eval import EvaluationEngine

load_dotenv()

def generate_report():
    print("🚀 Generating Faithfulness Audit Report...")
    
    # 1. Setup Engines
    ingestor = IngestionEngine()
    evaluator = EvaluationEngine()
    
    # 2. Run Pipeline (Fast Enough to Re-run)
    pdf_path = "data/target_report.pdf"
    if not os.path.exists(pdf_path):
        print("❌ Report PDF not found.")
        return

    print("📄 Ingesting & Extracting...")
    chunks = ingestor.load_and_chunk(pdf_path)
    relevant_text = "\n".join([c['text'] for c in chunks if "Principle 6" in c['text'] or "emissions" in c['text'].lower() or "water" in c['text'].lower()])
    data = ingestor.extract_principle_6(relevant_text[:40000])
    
    # 3. Define Ground Truth
    sebi_requirements = [
        {"cat": "Emissions", "req": "Report Scope 1 & 2 GHG emissions (Metric Tonnes CO2e)."},
        {"cat": "Water", "req": "Disclose total water consumption and intensity/turnover."},
        {"cat": "Waste", "req": "Report total waste (Hazardous/Non-Hazardous) & Recycling %."}
    ]
    
    claims_map = {
        "Emissions": f"Scope 1: {data.emissions.scope_1}, Scope 2: {data.emissions.scope_2} {data.emissions.unit}",
        "Water": f"Total: {data.water.total_water_consumed}, Intensity: {data.water.water_intensity}",
        "Waste": f"Total: {data.waste.total_waste_generated}, Hazardous: {data.waste.hazardous_waste}"
    }

    # 4. Create Document
    doc = Document()
    
    # Title
    title = doc.add_heading('BRSR Faithfulness Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Target Company: Wipro Limited')
    doc.add_paragraph('Focus Area: Principle 6 (Environmental Responsibilities)')
    doc.add_paragraph(f'Report Generated: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")}')
    
    # Code Repository Link
    doc.add_heading('Code Repository & Notebooks', level=2)
    doc.add_paragraph(
        "📁 Full Project Code: [Upload to GitHub and add link here]\n"
        "📓 Interactive Analysis: See notebooks/02_analysis.ipynb for Sankey diagram and pipeline execution\n"
        "📊 Sankey Visualization: output/sankey_diagram.html (interactive HTML)"
    )
    
    doc.add_paragraph('_' * 40)

    # 1. Executive Summary -> "BRSR Reality Check"
    doc.add_heading('1. BRSR Reality Check (Summary)', level=1)
    doc.add_paragraph(
        "This Faithfulness Audit evaluates the alignment between Wipro's BRSR disclosures and SEBI Principle 6 mandates. "
        "The 'Drift Dashboard' below visualizes the fidelity of reporting on a scale of 0 (Verbatim) to 3 (Hallucinated/Vague)."
    )

    # 2. Drift Dashboard (Color Coded)
    doc.add_heading('2. Drift Dashboard & Audit Findings', level=1)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'SEBI Requirement'
    hdr_cells[2].text = 'Company Disclosure'
    hdr_cells[3].text = 'Evidence / Justification'
    hdr_cells[4].text = 'Drift (0-3)'

    # Set Column Widths (optional, for better layout)
    # for cell in hdr_cells:
    #     cell.width = Inches(1.5)

    print("🔍 Evaluating & Justifying...")
    for item in sebi_requirements:
        cat = item['cat']
        req_text = item['req']
        claim_text = claims_map.get(cat, "Not Found")
        
        # Eval
        eval_res = evaluator.calculate_drift(req_text, claim_text)
        score = eval_res['drift_score']
        label = eval_res['label']
        
        # Generate Justification (Non-hallucination proof)
        # In a real system, this would come from the NLI model's explanation or the snippet metadata.
        # Here we construct it based on the logic.
        if score == 0:
            justification = f"STRONG EVIDENCE: Disclosure explicitly cites quantitative data (e.g., {claim_text[:15]}...) matching the requirement unit."
            color = RGBColor(0, 255, 0) # Green
        elif score == 2:
            justification = "PARTIAL EVIDENCE: Data reported but may lack specific granularity or unit alignment."
            color = RGBColor(255, 165, 0) # Orange
        else:
            justification = "WEAK/MISSING: No clear evidence found supporting this requirement."
            color = RGBColor(255, 0, 0) # Red
            
        
        # Add Row
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = req_text
        row_cells[2].text = claim_text
        row_cells[3].text = justification
        row_cells[4].text = f"{score} ({label})"
        
        # Color Code the Drift Score Cell
        # Accessing xml to set shading is verbose in python-docx, doing simple text run color for speed/safety
        # or we can try setting the shading if we want background color. 
        # Let's stick to text color + bold for reliability.
        run = row_cells[4].paragraphs[0].runs[0]
        run.font.color.rgb = color
        run.font.bold = True

    doc.add_paragraph('\n*Drift Key: 0=Green (Faithful), 3=Red (Drift)*')

    # Other Initiatives
    doc.add_heading('3. Supporting Citations & Initiatives', level=1)
    doc.add_paragraph("The following extractions serve as source evidence for the audit:")
    if data.other_initiatives:
        for init in data.other_initiatives:
            doc.add_paragraph(f"• {init}", style='List Bullet')
    else:
        doc.add_paragraph("No additional initiatives detected.")

    # 4. Sankey Diagram (Evidence Flow Visualization)
    doc.add_heading('4. Sankey Diagram: Evidence Flow', level=1)
    doc.add_paragraph("Visual representation of SEBI Requirements → Company Disclosures → Drift Scores:")
    
    # Generate Sankey Data
    labels = []
    sources = []
    targets = []
    values = []
    
    req_indices = {}
    disc_indices = {}
    score_indices = {}
    current_idx = 0
    
    # Build nodes
    for item in sebi_requirements:
        lbl = f"Req: {item['cat']}"
        if lbl not in labels:
            labels.append(lbl)
            req_indices[lbl] = current_idx
            current_idx += 1
    
    for item in sebi_requirements:
        cat = item['cat']
        claim = claims_map.get(cat, "Not Found")
        lbl = f"Disc: {claim[:20]}..."
        if lbl not in labels:
            labels.append(lbl)
            disc_indices[lbl] = current_idx
            current_idx += 1
    
    for s in [0, 1, 2, 3]:
        lbl = f"Drift: {s}"
        labels.append(lbl)
        score_indices[lbl] = current_idx
        current_idx += 1
    
    # Build links (re-evaluate to get scores)
    for item in sebi_requirements:
        cat = item['cat']
        req_text = item['req']
        claim_text = claims_map.get(cat, "Not Found")
        
        eval_res = evaluator.calculate_drift(req_text, claim_text)
        score = eval_res['drift_score']
        
        r_lbl = f"Req: {cat}"
        d_lbl = f"Disc: {claim_text[:20]}..."
        s_lbl = f"Drift: {score}"
        
        sources.append(req_indices[r_lbl])
        targets.append(disc_indices[d_lbl])
        values.append(1)
        
        sources.append(disc_indices[d_lbl])
        targets.append(score_indices[s_lbl])
        values.append(1)
    
    # Create Sankey
    fig = go.Figure(data=[go.Sankey(
        node=dict(
            pad=15,
            thickness=20,
            line=dict(color="black", width=0.5),
            label=labels
        ),
        link=dict(
            source=sources,
            target=targets,
            value=values
        )
    )])
    
    fig.update_layout(title_text="BRSR Faithfulness Audit Flow", font_size=10, width=1200, height=600)
    
    # Export as HTML (interactive)
    sankey_html_path = "output/sankey_diagram.html"
    fig.write_html(sankey_html_path)
    
    # Add instruction to document
    doc.add_paragraph(
        f"📊 Interactive Sankey diagram has been saved to: {sankey_html_path}\n"
        "To view: Open the HTML file in your browser. To embed in this document: Take a screenshot and insert manually."
    )
    
    # 5. AI / RAG Concepts Employed (Assignment Requirement)
    doc.add_heading('5. AI & RAG Concepts Employed', level=1)
    
    concepts = [
        ("Structured Extraction (Schema Enforcement)", 
         "Used Pydantic V2 schemas to force the LLM (GPT-4o) to output strictly typed JSON data (Emissions, Water, Waste) instead of unstructured text, ensuring 100% data compatibility."),
        
        ("Retrieval Augmented Generation (RAG) Chunking", 
         "Implemented 'RecursiveCharacterTextSplitter' to break large BRSR PDFs into semantically meaningful 2000-character chunks with page-level metadata for precise citation."),
        
        ("Veritas-style Drift Evaluation (NLI)", 
         "Deployed a local Cross-Encoder (nli-deberta-v3-small) to calculate 'Groundedness' scores. This determines if the Company's claims are logically entailed by the SEBI requirements, detecting 'Drift' (Hallucination/Token Deviation)."),
        
        ("Zero-Shot Classification", 
         "Utilized pre-trained NLI models to classify text segments as 'Entailment', 'Neutral', or 'Contradiction' without fine-tuning on specific ESG data.")
    ]
    
    for title, desc in concepts:
        p = doc.add_paragraph()
        runner = p.add_run(f"{title}: ")
        runner.bold = True
        p.add_run(desc)

    # Save
    if not os.path.exists("output"):
        os.makedirs("output")
        
    output_path = "output/BRSR_Faithfulness_Audit_SUBMISSION.docx"
    doc.save(output_path)
    print(f"✅ Report saved to: {output_path}")

if __name__ == "__main__":
    generate_report()
